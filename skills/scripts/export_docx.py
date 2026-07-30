import sys
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

def process_markdown_styles(paragraph, text):
    """Parses text for bold (** or __) and italics (* or _) and adds appropriate runs."""
    pattern = r'(\*\*(?P<bold1>.*?)\*\*|__(?P<bold2>.*?)__|_(?P<italic1>.*?)_|\*(?P<italic2>.*?)\*)'
    
    last_idx = 0
    for match in re.finditer(pattern, text):
        if match.start() > last_idx:
            paragraph.add_run(text[last_idx:match.start()])
        
        m_bold = match.group('bold1') or match.group('bold2')
        m_italic = match.group('italic1') or match.group('italic2')
        
        if m_bold:
            run = paragraph.add_run(m_bold)
            run.bold = True
        elif m_italic:
            run = paragraph.add_run(m_italic)
            run.italic = True
        
        last_idx = match.end()
    
    if last_idx < len(text):
        paragraph.add_run(text[last_idx:])

def set_cell_padding(cell, top=0, start=0, bottom=0, end=0):
    """Set cell margins (padding) in EMUs."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:start w:w="{start}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:end w:w="{end}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def style_table(table):
    """Apply clean, simple formatting to a completed table."""
    CELL_PAD = 80  # ~1.4mm in twips

    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_padding(cell, top=CELL_PAD, start=CELL_PAD, bottom=CELL_PAD, end=CELL_PAD)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
                    if row_idx == 0:
                        run.bold = True

def export_to_docx(markdown_path, docx_path):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Arial'
    h1_style.font.size = Pt(20)
    h1_style.font.bold = False
    
    h2_style = doc.styles['Heading 2']
    h2_style.font.name = 'Arial'
    h2_style.font.size = Pt(14)
    
    h3_style = doc.styles['Heading 3']
    h3_style.font.name = 'Arial'
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True
    
    with open(markdown_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    current_table = None

    for line in lines:
        line_stripped = line.rstrip()
        if not line_stripped:
            if in_table and current_table:
                style_table(current_table)
            in_table = False
            current_table = None
            continue
            
        leading_spaces = len(line) - len(line.lstrip())
        content = line.strip()

        if content.startswith('|') and content.endswith('|'):
            cells = [c.strip() for c in content.strip('|').split('|')]
            
            if all(all(char in '-: ' for char in c) for c in cells) and len(cells) > 0 and any('-' in c for c in cells):
                continue
                
            if not in_table:
                in_table = True
                current_table = doc.add_table(rows=1, cols=len(cells))
                current_table.style = 'Table Grid'
                hdr_cells = current_table.rows[0].cells
                for i, cell_text in enumerate(cells):
                    if i < len(hdr_cells):
                        p = hdr_cells[i].paragraphs[0]
                        process_markdown_styles(p, cell_text)
                        for run in p.runs:
                            run.bold = True
            else:
                row_cells = current_table.add_row().cells
                for i, cell_text in enumerate(cells):
                    if i < len(row_cells):
                        p = row_cells[i].paragraphs[0]
                        process_markdown_styles(p, cell_text)
            continue
        else:
            in_table = False
            current_table = None

        if content.startswith('# '):
            p = doc.add_heading(level=1)
            process_markdown_styles(p, content[2:])
        elif content.startswith('## '):
            p = doc.add_heading(level=2)
            process_markdown_styles(p, content[3:])
        elif content.startswith('### '):
            p = doc.add_heading(level=3)
            process_markdown_styles(p, content[4:])
        elif content.startswith('>'):
            quote_text = content[1:].strip()
            try:
                p = doc.add_paragraph(style='Quote')
            except KeyError:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
            process_markdown_styles(p, quote_text)
            for run in p.runs:
                run.italic = True
        elif content.startswith('- ') or content.startswith('* '):
            level = leading_spaces // 2
            style_name = 'List Bullet' if level == 0 else f'List Bullet {level + 1}'
            try:
                p = doc.add_paragraph(style=style_name)
            except KeyError:
                p = doc.add_paragraph(style='List Bullet')
            process_markdown_styles(p, content[2:])
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            process_markdown_styles(p, content)

    if in_table and current_table:
        style_table(current_table)

    doc.save(docx_path)
    print(f"Successfully created {docx_path}")

if __name__ == "__main__":
    import os
    if len(sys.argv) == 1:
        input_path = "output/memo.md"
        output_path = "output/memo.docx"
    elif len(sys.argv) == 2:
        input_path = sys.argv[1]
        base, _ = os.path.splitext(input_path)
        output_path = base + ".docx"
    elif len(sys.argv) == 3:
        input_path = sys.argv[1]
        output_path = sys.argv[2]
    else:
        print("Usage: python export_docx.py [input.md] [output.docx]")
        sys.exit(1)
    
    export_to_docx(input_path, output_path)
